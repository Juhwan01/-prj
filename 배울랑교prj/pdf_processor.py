# pdf_processor.py

from dotenv import load_dotenv
load_dotenv()

from typing import List, Dict
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from rag.pdf import PDFRetrievalChain
from typing import TypedDict, List, Dict
from langchain_upstage import UpstageGroundednessCheck
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from rag.utils import *
from langchain import hub
from langgraph.graph import END, StateGraph
from langgraph.checkpoint.memory import MemorySaver
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from langchain_core.runnables import RunnableConfig

class GraphState(TypedDict):
    question: str
    context: str
    answer: str
    relevance: str
    text: str
    qa_pairs: List[Dict[str, str]]

class QAPair(BaseModel):
    question: str = Field(description="The question from the original text")
    answer: str = Field(description="The answer to the question from the original text")

class QAList(BaseModel):
    qa_pairs: List[QAPair] = Field(description="List of question-answer pairs")

def extract_qa_pairs(answer: str) -> List[Dict[str, str]]:
    model = ChatOpenAI(temperature=0, model="gpt-4-0613")
    prompt = ChatPromptTemplate.from_template(
        """Extract all relevant question-answer pairs from the following text. 
        Format each pair as 'Q: [question] A: [answer]', with each pair on a new line.
        Make sure to capture all important information from the text in the form of questions and answers.

        Text: {answer}

        Question-Answer Pairs:"""
    )
    
    chain = prompt | model
    
    try:
        response = chain.invoke({"answer": answer})
        
        qa_pairs = response.content.split('\n')
        
        processed_pairs = []
        current_question = ""
        current_answer = ""
        
        for line in qa_pairs:
            line = line.strip()
            if line.startswith("Q:"):
                if current_question and current_answer:
                    processed_pairs.append({"question": current_question, "answer": current_answer})
                current_question = line[2:].strip()
                current_answer = ""
            elif line.startswith("A:"):
                current_answer += line[2:].strip() + " "
            else:
                current_answer += line + " "
        
        if current_question and current_answer:
            processed_pairs.append({"question": current_question, "answer": current_answer.strip()})
        
        return processed_pairs
    except Exception as e:
        print(f"Error in extract_qa_pairs: {str(e)}")
        print(f"Raw response: {response.content if 'response' in locals() else 'No response generated'}")
        return []

def create_qa_word_document(state: GraphState) -> None:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    doc.add_heading('Question and Answer Pairs', level=1)

    for i, pair in enumerate(state["qa_pairs"], 1):
        question_para = doc.add_paragraph()
        question_run = question_para.add_run(f"Q{i}: {pair['question']}")
        question_run.bold = True
        
        answer_para = doc.add_paragraph()
        answer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        answer_run = answer_para.add_run(f"A: {pair['answer']}")
        
        doc.add_paragraph()

    filename = "qa_output.docx"
    doc.save(filename)
    print(f"Word document '{filename}' has been created successfully.")

def setup_workflow(file_path: str, num_questions: int):
    pdf = PDFRetrievalChain([file_path]).create_chain()
    pdf_retriever = pdf.retriever
    pdf_chain = pdf.chain

    upstage_ground_checker = UpstageGroundednessCheck()

    def qMaker(state):
        qchain = qChain()
        response = qchain.invoke({"num_questions": num_questions, "context_str": state["text"]})
        print(num_questions)
        return GraphState(question=response)

    def llm_answer(state: GraphState) -> GraphState:
        question = state["question"]
        context = state["context"]
        
        model = ChatOpenAI(temperature=0, model="gpt-4-0613")
        prompt = hub.pull("lings/answer-generator")
        chain = prompt | model | StrOutputParser()
        response = chain.invoke({"question": question, "examContext": context})
        
        print(state["question"])
        print(response)
        return GraphState(answer=response)

    def a_relevance_check(state: GraphState) -> GraphState:
        response = upstage_ground_checker.run({"context": state["context"], "answer": state["answer"]})
        qa_pairs = extract_qa_pairs(state["answer"])
        return GraphState(
            relevance=response, 
            question=state["question"], 
            answer=state["answer"], 
            qa_pairs=qa_pairs
        )

    def retrieve_document(state: GraphState) -> GraphState:
        retrieved_docs = pdf_retriever.invoke(state["question"])
        retrieved_docs = format_docs(retrieved_docs)
        return GraphState(context=retrieved_docs)

    def docChecker(state: GraphState) -> GraphState:
        question = state["question"]
        context = state["context"]
        response = pdf_chain.invoke({"question": question, "context": context})
        return GraphState(relevance=response)

    def q_relevance_check(state: GraphState) -> GraphState:
        response = upstage_ground_checker.run({"context": state["text"], "answer": state["question"]})
        return GraphState(relevance=response, question=state["question"])

    def is_relevant(state: GraphState) -> GraphState:
        return state["relevance"]

    workflow = StateGraph(GraphState)

    workflow.add_node("retrieve", retrieve_document)
    workflow.add_node("document", create_qa_word_document)
    workflow.add_node("llm_answer", llm_answer)
    workflow.add_node("a_relevance_check", a_relevance_check)
    workflow.add_node("q_relevance_check", q_relevance_check)
    workflow.add_node("doccheck", docChecker)
    workflow.add_node("qmaker", qMaker)

    workflow.add_edge("retrieve", "doccheck")
    workflow.add_edge("llm_answer", "a_relevance_check")
    workflow.add_edge("qmaker", "q_relevance_check")
    workflow.add_edge("document", END)

    workflow.add_conditional_edges(
        "doccheck",
        is_relevant,
        {
            "yes": "llm_answer",
            "no": "retrieve",
        },
    )

    workflow.add_conditional_edges(
        "q_relevance_check",
        is_relevant,
        {
            "grounded": "retrieve",
            "notGrounded": "qmaker",
            "notSure": "qmaker",
        },
    )

    workflow.add_conditional_edges(
        "a_relevance_check",
        is_relevant,
        {
            "grounded": "document",
            "notGrounded": "retrieve",
            "notSure": "retrieve",
        },
    )

    workflow.set_entry_point("qmaker")

    return workflow

def process_pdf(file_path: str, num_questions: int):
    workflow = setup_workflow(file_path, num_questions)
    memory = MemorySaver()
    app = workflow.compile(checkpointer=memory)

    config = RunnableConfig(recursion_limit=100, configurable={"thread_id": "CORRECTIVE-SEARCH-RAG"})
    inputs = GraphState(text=load_doc(file_path))
    result = app.invoke(inputs, config=config)

    return result
