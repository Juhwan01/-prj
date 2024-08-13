from dotenv import load_dotenv
load_dotenv()

from rag.pdf import PDFRetrievalChain
from typing import TypedDict, List, Dict
from langchain_upstage import UpstageGroundednessCheck
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from rag.utils import *
from langchain import hub
from langgraph.graph import END, StateGraph
from langgraph.checkpoint.memory import MemorySaver
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt

# Existing code...
pdf = PDFRetrievalChain(["배울랑교prj/testcheck.pdf"]).create_chain()
pdf_retriever = pdf.retriever
pdf_chain = pdf.chain

class GraphState(TypedDict):
    question: str
    context: str
    answer: str
    relevance: str
    text: str
    qa_pairs: List[Dict[str, str]]  # Field for extracted Q&A pairs

upstage_ground_checker = UpstageGroundednessCheck()
qchain = qChain()

# Classes for Q&A extraction
class QAPair(BaseModel):
    question: str = Field(description="The question from the original text")
    answer: str = Field(description="The answer to the question from the original text")

class QAList(BaseModel):
    qa_pairs: List[QAPair] = Field(description="List of question-answer pairs")

# Function for Q&A extraction
def extract_qa_pairs(answer: str) -> List[Dict[str, str]]:
    model = ChatOpenAI(temperature=0, model="gpt-4-0613")
    prompt = ChatPromptTemplate.from_template(
        """Extract all relevant question-answer pairs from the following text. 
        Format each pair as 'Q: [question] A: [answer]', with each pair on a new line.
        Make sure to capture all important information from the text in the form of questions and answers.

        Text: {answer}

        Question-Answer Pairs:"""
    )
    
    chain = prompt | model
    
    try:
        response = chain.invoke({"answer": answer})
        
        # Split the response into individual QA pairs
        qa_pairs = response.content.split('\n')
        
        # Process each pair and create a list of dictionaries
        processed_pairs = []
        current_question = ""
        current_answer = ""
        
        for line in qa_pairs:
            line = line.strip()
            if line.startswith("Q:"):
                # If we have a previous QA pair, add it to the list
                if current_question and current_answer:
                    processed_pairs.append({"question": current_question, "answer": current_answer})
                # Start a new question
                current_question = line[2:].strip()
                current_answer = ""
            elif line.startswith("A:"):
                # Add to the current answer
                current_answer += line[2:].strip() + " "
            else:
                # Continuation of the previous answer
                current_answer += line + " "
        
        # Add the last QA pair if it exists
        if current_question and current_answer:
            processed_pairs.append({"question": current_question, "answer": current_answer.strip()})
        
        return processed_pairs
    except Exception as e:
        print(f"Error in extract_qa_pairs: {str(e)}")
        print(f"Raw response: {response.content if 'response' in locals() else 'No response generated'}")
        return []

# Modified llm_answer function (no Q&A extraction here)
def llm_answer(state: GraphState) -> GraphState:
    question = state["question"]
    context = state["context"]
    
    model = ChatOpenAI(temperature=0, model="gpt-4-0613")
    prompt = hub.pull("lings/answer-generator")
    chain = prompt | model | StrOutputParser()
    response = chain.invoke({"question": question, "examContext": context})
    
    print(state["question"])
    print(response)
    return GraphState(answer=response)

# Modified a_relevance_check function with Q&A extraction
def a_relevance_check(state: GraphState) -> GraphState:
    response = upstage_ground_checker.run({"context": state["context"], "answer": state["answer"]})
    qa_pairs = extract_qa_pairs(state["answer"])
    return GraphState(
        relevance=response, 
        question=state["question"], 
        answer=state["answer"], 
        qa_pairs=qa_pairs
    )

# Existing functions...
def retrieve_document(state: GraphState) -> GraphState:
    retrieved_docs = pdf_retriever.invoke(state["question"])
    retrieved_docs = format_docs(retrieved_docs)
    return GraphState(context=retrieved_docs)

def docChecker(state: GraphState) -> GraphState:
    question = state["question"]
    context = state["context"]
    response = pdf_chain.invoke({"question": question, "context": context})
    return GraphState(relevance=response)

def qMaker(state):
    response = qchain.invoke({"num_questions": 3, "context_str": state["text"]})
    return GraphState(question=response)

def q_relevance_check(state: GraphState) -> GraphState:
    response = upstage_ground_checker.run({"context": state["text"], "answer": state["question"]})
    return GraphState(relevance=response, question=state["question"])

def is_relevant(state: GraphState) -> GraphState:
    return state["relevance"]

# New function to save questions and answers to a Word document
def save_qa_to_word(qa_pairs: List[Dict[str, str]]):
    doc = Document()
    
    # Set the font for the entire document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for pair in qa_pairs:
        # Add question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run('질문: ')
        q_run.bold = True
        q_para.add_run(pair['question'])

        # Add answer
        a_para = doc.add_paragraph()
        a_run = a_para.add_run('답변: ')
        a_run.bold = True
        a_para.add_run(pair['answer'])

        # Add a blank line between Q&A pairs
        doc.add_paragraph()

    # Save the document
    doc.save('질문과_답변.docx')

# Workflow setup
workflow = StateGraph(GraphState)

workflow.add_node("retrieve", retrieve_document)
workflow.add_node("llm_answer", llm_answer)
workflow.add_node("a_relevance_check", a_relevance_check)
workflow.add_node("q_relevance_check", q_relevance_check)
workflow.add_node("doccheck", docChecker)
workflow.add_node("qmaker", qMaker)

workflow.add_edge("retrieve", "doccheck")
workflow.add_edge("llm_answer", "a_relevance_check")
workflow.add_edge("qmaker", "q_relevance_check")

workflow.add_conditional_edges(
    "doccheck",
    is_relevant,
    {
        "yes": "llm_answer",
        "no": "retrieve",
    },
)

workflow.add_conditional_edges(
    "q_relevance_check",
    is_relevant,
    {
        "grounded": "retrieve",
        "notGrounded": "qmaker",
        "notSure": "qmaker",
    },
)

workflow.add_conditional_edges(
    "a_relevance_check",
    is_relevant,
    {
        "grounded": END,
        "notGrounded": "retrieve",
        "notSure": "retrieve",
    },
)

workflow.set_entry_point("qmaker")

memory = MemorySaver()
app = workflow.compile(checkpointer=memory)

# Usage
from langchain_core.runnables import RunnableConfig

config = RunnableConfig(recursion_limit=100, configurable={"thread_id": "CORRECTIVE-SEARCH-RAG"})
inputs = GraphState(text=load_doc("배울랑교prj/testcheck.pdf"))
result = app.invoke(inputs, config=config)

# Print extracted Q&A pairs and save to a Word document
if "qa_pairs" in result:
    print("Extracted Q&A Pairs:")
    for pair in result["qa_pairs"]:
        print(f"Q: {pair['question']}")
        print(f"A: {pair['answer']}")
        print(pair)
    
    # Save Q&A pairs to a Word document
    save_qa_to_word(result["qa_pairs"])
    print("Questions and answers have been saved to '질문과_답변.docx'.")